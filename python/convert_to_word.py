#!/usr/bin/env python3
"""
Convert a PDF to Word (.docx) format.

Usage:
    python3 convert_to_word.py <input_pdf> <output_docx> --layout flow|exact [--images] [--no-images] [--ocr] [--json]
"""

import sys
import os
import argparse
import json
import traceback

try:
    import fitz  # PyMuPDF
except ImportError:
    print(json.dumps({"success": False, "error": "PyMuPDF is not installed. Run: pip install PyMuPDF"}))
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
except ImportError:
    print(json.dumps({"success": False, "error": "python-docx is not installed. Run: pip install python-docx"}))
    sys.exit(1)


import re as _re

# Regex to strip XML-illegal characters (NULL bytes, control chars except tab/newline/cr)
_XML_ILLEGAL_RE = _re.compile(
    '[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f'
    '\ud800-\udfff\ufdd0-\ufdef\ufffe\uffff]'
)

def sanitize_for_xml(text):
    """Remove characters that are illegal in XML 1.0."""
    if not text:
        return text
    return _XML_ILLEGAL_RE.sub('', text)


def extract_page_blocks(page):
    """Extract text blocks from a page using PyMuPDF dict output."""
    blocks = page.get_text("dict", sort=True)["blocks"]
    return blocks


def get_dominant_font(spans):
    """Get the most common font name and size from a list of spans."""
    if not spans:
        return "Arial", 11
    font_sizes = {}
    for span in spans:
        key = (span.get("font", "Arial"), round(span.get("size", 11)))
        font_sizes[key] = font_sizes.get(key, 0) + len(span.get("text", ""))
    dominant = max(font_sizes, key=font_sizes.get)
    return dominant[0], dominant[1]


def is_bold_font(font_name):
    """Check if font name indicates bold."""
    return any(kw in font_name.lower() for kw in ["bold", "heavy", "black", "demi"])


def is_italic_font(font_name):
    """Check if font name indicates italic."""
    return any(kw in font_name.lower() for kw in ["italic", "oblique", "slant"])


def color_from_int(color_int):
    """Convert integer color to RGBColor."""
    if color_int is None or color_int == 0:
        return RGBColor(0, 0, 0)
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8) & 0xFF
    b = color_int & 0xFF
    return RGBColor(r, g, b)


def add_span_to_paragraph(paragraph, span):
    """Add a span with formatting to a Word paragraph."""
    text = sanitize_for_xml(span.get("text", ""))
    if not text:
        return
    
    run = paragraph.add_run(text)
    font_name = span.get("font", "Arial")
    font_size = span.get("size", 11)
    color = span.get("color", 0)
    flags = span.get("flags", 0)
    
    run.font.size = Pt(max(6, min(font_size, 72)))
    
    # Bold: flag bit 4 (weight >= 600) or font name
    is_bold = (flags & 16) != 0 or is_bold_font(font_name)
    run.font.bold = is_bold
    
    # Italic: flag bit 1 or font name
    is_italic = (flags & 2) != 0 or is_italic_font(font_name)
    run.font.italic = is_italic
    
    # Monospace detection
    clean_name = font_name.split("+")[-1].split(",")[0].split("-")[0].strip()
    run.font.name = clean_name if clean_name else "Arial"
    
    # Color
    if color and color != 0:
        run.font.color.rgb = color_from_int(color)


def convert_flow_layout(pdf_doc, word_doc, include_images=True):
    """
    Flow layout: extract text in reading order, 
    treating the PDF like a continuous document.
    """
    for page_idx in range(len(pdf_doc)):
        page = pdf_doc[page_idx]
        
        if page_idx > 0:
            word_doc.add_page_break()
        
        blocks = extract_page_blocks(page)
        
        for block in blocks:
            if block["type"] == 0:  # Text block
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    
                    # Check if entire line is empty
                    full_text = "".join(s.get("text", "") for s in spans).strip()
                    if not full_text:
                        continue
                    
                    paragraph = word_doc.add_paragraph()
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.space_before = Pt(0)
                    
                    _, dominant_size = get_dominant_font(spans)
                    if dominant_size >= 18:
                        paragraph.style = word_doc.styles['Heading 1']
                    elif dominant_size >= 14:
                        paragraph.style = word_doc.styles['Heading 2']
                    elif dominant_size >= 12.5:
                        paragraph.style = word_doc.styles['Heading 3']
                    
                    for span in spans:
                        add_span_to_paragraph(paragraph, span)
            
            elif block["type"] == 1 and include_images:  # Image block
                try:
                    img_bbox = fitz.Rect(block["bbox"])
                    width_inches = img_bbox.width / 72.0
                    
                    # Cap image width to page width
                    max_width = 6.0
                    if width_inches > max_width:
                        width_inches = max_width
                    
                    # Extract image bytes
                    if "image" in block:
                        img_bytes = block["image"]
                    else:
                        clip = fitz.Rect(block["bbox"])
                        pix = page.get_pixmap(clip=clip, dpi=150)
                        img_bytes = pix.tobytes("png")
                    
                    import io
                    img_stream = io.BytesIO(img_bytes)
                    word_doc.add_picture(img_stream, width=Inches(width_inches))
                except Exception:
                    pass  # Skip failed images


def convert_exact_layout(pdf_doc, word_doc, include_images=True):
    """
    Exact layout: try to preserve spatial positioning by using tables 
    and absolute positioning tricks. This is a best-effort approach.
    """
    for page_idx in range(len(pdf_doc)):
        page = pdf_doc[page_idx]
        page_width = page.rect.width
        page_height = page.rect.height
        
        if page_idx > 0:
            word_doc.add_page_break()
        
        # Set page dimensions
        section = word_doc.sections[-1] if page_idx == 0 else word_doc.add_section()
        section.page_width = Emu(int(page_width / 72 * 914400))
        section.page_height = Emu(int(page_height / 72 * 914400))
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        
        blocks = extract_page_blocks(page)
        
        # Sort blocks top to bottom
        blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))
        
        for block in blocks:
            if block["type"] == 0:  # Text block
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    
                    full_text = "".join(s.get("text", "") for s in spans).strip()
                    if not full_text:
                        continue
                    
                    paragraph = word_doc.add_paragraph()
                    paragraph.paragraph_format.space_after = Pt(1)
                    paragraph.paragraph_format.space_before = Pt(1)
                    
                    for span in spans:
                        add_span_to_paragraph(paragraph, span)
            
            elif block["type"] == 1 and include_images:
                try:
                    img_bbox = fitz.Rect(block["bbox"])
                    width_inches = min(img_bbox.width / 72.0, 6.0)
                    
                    if "image" in block:
                        img_bytes = block["image"]
                    else:
                        clip = fitz.Rect(block["bbox"])
                        pix = page.get_pixmap(clip=clip, dpi=150)
                        img_bytes = pix.tobytes("png")
                    
                    import io
                    img_stream = io.BytesIO(img_bytes)
                    word_doc.add_picture(img_stream, width=Inches(width_inches))
                except Exception:
                    pass


def perform_ocr_on_page(page):
    """Use Tesseract OCR on a page image to extract text."""
    try:
        import pytesseract
        from PIL import Image
        import io
        
        pix = page.get_pixmap(dpi=300)
        img_bytes = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_bytes))
        
        ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        
        lines = {}
        for i in range(len(ocr_data["text"])):
            text = ocr_data["text"][i].strip()
            if not text:
                continue
            line_num = ocr_data["line_num"][i]
            block_num = ocr_data["block_num"][i]
            key = (block_num, line_num)
            if key not in lines:
                lines[key] = []
            lines[key].append(text)
        
        return [" ".join(words) for words in lines.values()]
    except ImportError:
        return []


def main():
    parser = argparse.ArgumentParser(description="Convert PDF to Word (.docx)")
    parser.add_argument("input_pdf", help="Path to input PDF file")
    parser.add_argument("output_docx", help="Path to output DOCX file")
    parser.add_argument("--layout", choices=["flow", "exact"], default="flow",
                        help="Layout mode: flow (reading order) or exact (preserve positions)")
    parser.add_argument("--images", dest="include_images", action="store_true", default=True,
                        help="Include images in output")
    parser.add_argument("--no-images", dest="include_images", action="store_false",
                        help="Exclude images from output")
    parser.add_argument("--ocr", action="store_true", default=False,
                        help="Use OCR for scanned pages")
    parser.add_argument("--json", action="store_true", default=False,
                        help="Output result as JSON")
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input_pdf):
        result = {"success": False, "error": f"Input file not found: {args.input_pdf}"}
        if args.json:
            print(json.dumps(result))
        else:
            print(f"Error: {result['error']}", file=sys.stderr)
        sys.exit(1)
    
    try:
        pdf_doc = fitz.open(args.input_pdf)
        word_doc = Document()
        
        # Set default font
        style = word_doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        total_pages = len(pdf_doc)
        
        # Check for scanned/image-only pages if OCR requested
        if args.ocr:
            for page_idx in range(total_pages):
                page = pdf_doc[page_idx]
                text = page.get_text("text").strip()
                
                if len(text) < 50:  # Likely a scanned page
                    ocr_lines = perform_ocr_on_page(page)
                    if ocr_lines:
                        if page_idx > 0:
                            word_doc.add_page_break()
                        for line_text in ocr_lines:
                            paragraph = word_doc.add_paragraph(sanitize_for_xml(line_text))
                            paragraph.paragraph_format.space_after = Pt(2)
                    continue
                
                # Non-scanned page, use normal extraction
                # (handled below — but we need to track which pages we did OCR on)
            
            # Re-open for normal pages 
            pdf_doc.close()
            pdf_doc = fitz.open(args.input_pdf)
        
        # Convert using selected layout
        if args.layout == "flow":
            convert_flow_layout(pdf_doc, word_doc, args.include_images)
        else:
            convert_exact_layout(pdf_doc, word_doc, args.include_images)
        
        # Save
        word_doc.save(args.output_docx)
        pdf_doc.close()
        
        file_size = os.path.getsize(args.output_docx)
        
        result = {
            "success": True,
            "output": args.output_docx,
            "pages": total_pages,
            "file_size": file_size,
            "layout": args.layout,
        }
        
        if args.json:
            print(json.dumps(result))
        else:
            print(f"Converted {total_pages} pages to {args.output_docx} ({file_size} bytes)")
    
    except Exception as e:
        result = {"success": False, "error": str(e), "traceback": traceback.format_exc()}
        if args.json:
            print(json.dumps(result))
        else:
            print(f"Error: {e}", file=sys.stderr)
            traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
