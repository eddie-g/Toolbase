import json
import os
import subprocess
import tempfile
import unittest
import io
from pathlib import Path
from zipfile import ZipFile

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
from PIL import Image


ROOT = Path(__file__).resolve().parents[3]
PYTHON = ROOT / "python" / "venv" / "bin" / "python"
CONVERTER = ROOT / "python" / "pdf-editor" / "convert_to_word.py"
FIXTURE = ROOT / "tests" / "pdfjs" / "test_pdf_1.pdf"
DRYLAB_FIXTURE = ROOT / "tests" / "OverlayEditor" / "drylab.pdf"


class WordLayoutConversionTest(unittest.TestCase):
    def convert(self, source, output, extra_args=None):
        extra_args = list(extra_args or [])
        result = subprocess.run(
            [
                str(PYTHON),
                str(CONVERTER),
                str(source),
                str(output),
                "--layout",
                "exact",
                "--images",
                *extra_args,
                "--json",
            ],
            cwd=ROOT,
            env={**os.environ, "PYTHONDONTWRITEBYTECODE": "1"},
            capture_output=True,
            text=True,
            check=False,
        )
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        self.assertTrue(output.is_file())
        json_lines = []
        for line in result.stdout.splitlines():
            try:
                json_lines.append(json.loads(line))
            except json.JSONDecodeError:
                pass
        self.assertTrue(json_lines, result.stdout)
        self.assertTrue(json_lines[-1]["success"])
        return json_lines[-1]

    def test_exact_layout_preserves_tables_paragraph_spacing_and_alignment(self):
        with tempfile.TemporaryDirectory(prefix="toolbase_word_test_") as directory:
            output = Path(directory) / "layout.docx"
            result = self.convert(FIXTURE, output)
            self.assertEqual(result["engine"], "pdf2docx")

            document = Document(output)
            self.assertEqual(len(document.tables), 2)
            expected = [["1", "2", "3", "4"], ["5", "6", "7", "8"], ["9", "10", "11", "12"], ["13", "14", "15", "16"]]
            for table in document.tables:
                actual = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                self.assertEqual(actual, expected)

            paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
            first = next(paragraph for paragraph in paragraphs if paragraph.text.startswith("Lorem ipsum dolor sit amet"))
            second = next(paragraph for paragraph in paragraphs if paragraph.text.startswith("Aenean pulvinar euismod ligula"))
            after_table = next(paragraph for paragraph in paragraphs if paragraph.text.startswith("Fusce efficitur mi ex"))

            self.assertEqual(first.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
            self.assertGreater(float(first.paragraph_format.line_spacing), 1)
            self.assertGreater(second.paragraph_format.space_before.pt, 10)
            self.assertGreater(after_table.paragraph_format.space_before.pt, 70)

            standalone_numbers = {
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip().isdigit()
            }
            self.assertEqual(standalone_numbers, set())

            first_section = document.sections[0]
            self.assertAlmostEqual(first_section.left_margin.pt, 66.6, delta=2)
            self.assertAlmostEqual(first_section.right_margin.pt, 66.5, delta=2)

    def test_exact_layout_reconstructs_drylab_columns_and_optimizes_photos(self):
        with tempfile.TemporaryDirectory(prefix="toolbase_word_drylab_test_") as directory:
            output = Path(directory) / "drylab.docx"
            result = self.convert(DRYLAB_FIXTURE, output)

            self.assertEqual(result["engine"], "pdf2docx")
            self.assertGreaterEqual(result["optimized_images"], 2)
            self.assertLess(output.stat().st_size, 2_000_000)

            document = Document(output)
            self.assertGreaterEqual(len(document.sections), 6)
            self.assertGreaterEqual(len(document.inline_shapes), 3)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("Welcome to our first newsletter", text)
            self.assertIn("Annual General Meeting", text)

            with ZipFile(output) as archive:
                document_xml = etree.fromstring(archive.read("word/document.xml"))
            namespaces = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            }
            two_column_sections = document_xml.xpath(
                './/w:sectPr/w:cols[@w:num="2"]', namespaces=namespaces
            )
            self.assertGreaterEqual(len(two_column_sections), 3)

    def test_visual_fidelity_exact_preserves_rotated_page_paint(self):
        with tempfile.TemporaryDirectory(prefix="toolbase_word_visual_test_") as directory:
            source = Path(directory) / "edited-rotated.pdf"
            output = Path(directory) / "edited-rotated.docx"
            pdf = fitz.open()
            page = pdf.new_page(width=460, height=300)
            page.insert_text((40, 55), "VISIBLE EDITED TEXT", fontsize=18, color=(1, 0, 0))
            page.draw_line((30, 80), (420, 240), color=(1, 0.25, 0), width=8)
            page.set_rotation(90)
            pdf.save(source)
            pdf.close()

            result = self.convert(source, output, ["--visual-fidelity"])

            self.assertEqual(result["engine"], "toolbase_visual")
            self.assertTrue(result["visual_fidelity"])
            self.assertEqual(result["rendered_pages"], 1)
            document = Document(output)
            self.assertEqual(len(document.inline_shapes), 1)
            self.assertAlmostEqual(document.sections[0].page_width.pt, 300, delta=1)
            self.assertAlmostEqual(document.sections[0].page_height.pt, 460, delta=1)

            with ZipFile(output) as archive:
                image_name = next(name for name in archive.namelist() if name.startswith("word/media/"))
                rendered = Image.open(io.BytesIO(archive.read(image_name))).convert("RGB")
            red_pixels = sum(
                1
                for red, green, blue in rendered.getdata()
                if red > 180 and green < 120 and blue < 120
            )
            self.assertGreater(red_pixels, 100)


if __name__ == "__main__":
    unittest.main()
