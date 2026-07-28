#!/usr/bin/env python3
"""
Convert a PDF to Word (.docx) format.

Usage:
    python3 convert_to_word.py <input_pdf> <output_docx> --layout flow|exact [--images] [--no-images] [--ocr] [--json]
"""

import sys
import os
import argparse
import copy
import io
import json
import logging
import traceback
import tempfile
import zipfile
from pathlib import Path
from statistics import median

try:
    import fitz  # PyMuPDF
except ImportError:
    print(json.dumps({"success": False, "error": "PyMuPDF is not installed. Run: pip install PyMuPDF"}))
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Emu
    from docx.enum.text import WD_BREAK, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print(json.dumps({"success": False, "error": "python-docx is not installed. Run: pip install python-docx"}))
    sys.exit(1)

try:
    from pdf2docx import Converter as Pdf2DocxConverter
except ImportError:
    Pdf2DocxConverter = None


import re as _re

# Regex to strip XML-illegal characters (NULL bytes, control chars except tab/newline/cr)
_XML_ILLEGAL_RE = _re.compile(
    '[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f'
    '\ud800-\udfff\ufdd0-\ufdef\ufffe\uffff]'
)

def sanitize_for_xml(text):
    """Remove characters that are illegal in XML 1.0."""
    if not text:
        return text
    return _XML_ILLEGAL_RE.sub('', text)


def _optimize_docx_images(output_docx):
    """Replace oversized opaque PNGs with visually equivalent JPEGs."""
    try:
        from PIL import Image
    except ImportError:
        return {"optimized_images": 0, "bytes_saved": 0}

    docx_path = Path(output_docx)
    replacements = {}
    original_size = docx_path.stat().st_size
    original_mode = docx_path.stat().st_mode

    with zipfile.ZipFile(docx_path, "r") as archive:
        for name in archive.namelist():
            if not name.startswith("word/media/") or not name.lower().endswith(".png"):
                continue
            source_bytes = archive.read(name)
            try:
                image = Image.open(io.BytesIO(source_bytes))
                image.load()
            except Exception:
                continue

            if "A" in image.getbands():
                alpha_min, _ = image.getchannel("A").getextrema()
                if alpha_min < 255:
                    continue

            converted = io.BytesIO()
            image.convert("RGB").save(converted, format="JPEG", quality=90, optimize=True)
            converted_bytes = converted.getvalue()
            if len(converted_bytes) >= len(source_bytes) * 0.9:
                continue

            new_name = str(Path(name).with_suffix(".jpg"))
            replacements[name] = (new_name, converted_bytes)

    if not replacements:
        return {"optimized_images": 0, "bytes_saved": 0}

    with tempfile.NamedTemporaryFile(
        prefix=f".{docx_path.stem}_",
        suffix=".docx",
        dir=docx_path.parent,
        delete=False,
    ) as temporary:
        temporary_path = Path(temporary.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(
            temporary_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as destination:
            for info in source.infolist():
                data = source.read(info.filename)
                output_name = info.filename

                if info.filename in replacements:
                    output_name, data = replacements[info.filename]
                elif info.filename.endswith((".xml", ".rels")):
                    for old_name, (new_name, _) in replacements.items():
                        data = data.replace(
                            Path(old_name).name.encode("utf-8"),
                            Path(new_name).name.encode("utf-8"),
                        )
                    if info.filename == "[Content_Types].xml" and b'Extension="jpg"' not in data:
                        data = data.replace(
                            b"</Types>",
                            b'<Default Extension="jpg" ContentType="image/jpeg"/></Types>',
                        )

                output_info = copy.copy(info)
                output_info.filename = output_name
                destination.writestr(output_info, data)

        os.chmod(temporary_path, original_mode)
        os.replace(temporary_path, docx_path)
    finally:
        temporary_path.unlink(missing_ok=True)

    return {
        "optimized_images": len(replacements),
        "bytes_saved": max(0, original_size - docx_path.stat().st_size),
    }


def _convert_with_layout_engine(input_pdf, output_docx):
    """Use pdf2docx's column, table, and page-layout reconstruction engine."""
    if Pdf2DocxConverter is None:
        raise RuntimeError("pdf2docx is not installed")

    # Keep third-party progress output out of the API's JSON response stream.
    logging.getLogger("pdf2docx").setLevel(logging.ERROR)
    converter = Pdf2DocxConverter(input_pdf)
    try:
        converter.convert(
            output_docx,
            multi_processing=False,
            clip_image_res_ratio=2.0,
            raw_exceptions=True,
        )
    finally:
        converter.close()

    return _optimize_docx_images(output_docx)


def extract_page_blocks(page):
    """Extract text blocks from a page using PyMuPDF dict output."""
    blocks = page.get_text("dict", sort=True)["blocks"]
    return blocks


def get_dominant_font(spans):
    """Get the most common font name and size from a list of spans."""
    if not spans:
        return "Arial", 11
    font_sizes = {}
    for span in spans:
        key = (span.get("font", "Arial"), round(span.get("size", 11)))
        font_sizes[key] = font_sizes.get(key, 0) + len(span.get("text", ""))
    dominant = max(font_sizes, key=font_sizes.get)
    return dominant[0], dominant[1]


def is_bold_font(font_name):
    """Check if font name indicates bold."""
    return any(kw in font_name.lower() for kw in ["bold", "heavy", "black", "demi"])


def is_italic_font(font_name):
    """Check if font name indicates italic."""
    return any(kw in font_name.lower() for kw in ["italic", "oblique", "slant"])


def color_from_int(color_int):
    """Convert integer color to RGBColor."""
    if color_int is None or color_int == 0:
        return RGBColor(0, 0, 0)
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8) & 0xFF
    b = color_int & 0xFF
    return RGBColor(r, g, b)


def add_span_to_paragraph(paragraph, span):
    """Add a span with formatting to a Word paragraph."""
    text = sanitize_for_xml(span.get("text", ""))
    if not text:
        return
    
    run = paragraph.add_run(text)
    font_name = span.get("font", "Arial")
    font_size = span.get("size", 11)
    color = span.get("color", 0)
    flags = span.get("flags", 0)
    
    run.font.size = Pt(max(6, min(font_size, 72)))
    
    # Bold: flag bit 4 (weight >= 600) or font name
    is_bold = (flags & 16) != 0 or is_bold_font(font_name)
    run.font.bold = is_bold
    
    # Italic: flag bit 1 or font name
    is_italic = (flags & 2) != 0 or is_italic_font(font_name)
    run.font.italic = is_italic
    
    # Monospace detection
    clean_name = font_name.split("+")[-1].split(",")[0].split("-")[0].strip()
    run.font.name = clean_name if clean_name else "Arial"
    
    # Color
    if color and color != 0:
        run.font.color.rgb = color_from_int(color)

    return run


def _nonempty_lines(block):
    return [
        line for line in block.get("lines", [])
        if "".join(span.get("text", "") for span in line.get("spans", [])).strip()
    ]


def _line_is_inside_table(line, table_bboxes):
    bbox = line.get("bbox", (0, 0, 0, 0))
    center_x = (bbox[0] + bbox[2]) / 2
    center_y = (bbox[1] + bbox[3]) / 2
    return any(
        table_bbox[0] - 1 <= center_x <= table_bbox[2] + 1
        and table_bbox[1] - 1 <= center_y <= table_bbox[3] + 1
        for table_bbox in table_bboxes
    )


def _rect_overlap_ratio(first, second):
    left = max(first[0], second[0])
    top = max(first[1], second[1])
    right = min(first[2], second[2])
    bottom = min(first[3], second[3])
    intersection = max(0, right - left) * max(0, bottom - top)
    area = max(1, (first[2] - first[0]) * (first[3] - first[1]))
    return intersection / area


def _detect_page_tables(page):
    try:
        return list(page.find_tables().tables)
    except Exception:
        return []


def _build_page_events(page, include_images=True):
    tables = _detect_page_tables(page)
    table_bboxes = [tuple(table.bbox) for table in tables]
    events = [{"type": "table", "bbox": tuple(table.bbox), "table": table} for table in tables]

    for block in extract_page_blocks(page):
        block_bbox = tuple(block.get("bbox", (0, 0, 0, 0)))
        if block.get("type") == 0:
            lines = [
                line for line in _nonempty_lines(block)
                if not _line_is_inside_table(line, table_bboxes)
            ]
            if not lines:
                continue
            event_bbox = (
                min(line["bbox"][0] for line in lines),
                min(line["bbox"][1] for line in lines),
                max(line["bbox"][2] for line in lines),
                max(line["bbox"][3] for line in lines),
            )
            events.append({"type": "text", "bbox": event_bbox, "lines": lines})
        elif block.get("type") == 1 and include_images:
            if any(_rect_overlap_ratio(block_bbox, table_bbox) > 0.5 for table_bbox in table_bboxes):
                continue
            events.append({"type": "image", "bbox": block_bbox, "block": block})

    return sorted(events, key=lambda event: (event["bbox"][1], event["bbox"][0]))


def _page_content_bounds(page, events):
    if not events:
        return (36.0, 36.0, 36.0, 36.0)
    left = max(18.0, min(144.0, min(event["bbox"][0] for event in events)))
    right = max(18.0, min(144.0, page.rect.width - max(event["bbox"][2] for event in events)))
    top = max(18.0, min(144.0, min(event["bbox"][1] for event in events)))
    bottom = max(18.0, min(144.0, page.rect.height - max(event["bbox"][3] for event in events)))
    return (left, right, top, bottom)


def _source_line_spacing(lines):
    line_tops = [float(line["bbox"][1]) for line in lines]
    gaps = [line_tops[index] - line_tops[index - 1] for index in range(1, len(line_tops))]
    spans = [span for line in lines for span in line.get("spans", [])]
    _, font_size = get_dominant_font(spans)
    return max(float(font_size) * 1.05, median(gaps) if gaps else float(font_size) * 1.15)


def _add_text_event(word_doc, event, content_left, content_right, gap_before=0, exact=False):
    lines = event["lines"]
    paragraph = word_doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(max(0, gap_before))
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(max(0, event["bbox"][0] - content_left))
    paragraph.paragraph_format.right_indent = Pt(max(0, content_right - event["bbox"][2]))
    paragraph.paragraph_format.keep_together = True

    line_spacing = _source_line_spacing(lines)
    if exact:
        paragraph.paragraph_format.line_spacing = Pt(line_spacing)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    else:
        paragraph.paragraph_format.line_spacing = max(1.0, min(2.0, line_spacing / max(1, get_dominant_font([
            span for line in lines for span in line.get("spans", [])
        ])[1])))

    for line_index, line in enumerate(lines):
        if line_index > 0:
            if exact:
                paragraph.add_run().add_break(WD_BREAK.LINE)
            elif not paragraph.text.endswith((" ", "-")):
                paragraph.add_run(" ")
        for span in line.get("spans", []):
            add_span_to_paragraph(paragraph, span)
    return paragraph


def _set_cell_margins(cell, top=30, start=45, bottom=30, end=45):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _cell_source_lines(page, bbox):
    lines = []
    clip = fitz.Rect(bbox)
    for block in page.get_text("dict", clip=clip, sort=True).get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in _nonempty_lines(block):
            center_x = (line["bbox"][0] + line["bbox"][2]) / 2
            center_y = (line["bbox"][1] + line["bbox"][3]) / 2
            if clip.x0 - 1 <= center_x <= clip.x1 + 1 and clip.y0 - 1 <= center_y <= clip.y1 + 1:
                lines.append(line)
    return lines


def _populate_table_cell(page, cell, bbox, fallback_text=""):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    lines = _cell_source_lines(page, bbox) if bbox else []
    if lines:
        paragraph.paragraph_format.line_spacing = Pt(_source_line_spacing(lines))
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        for line_index, line in enumerate(lines):
            if line_index > 0:
                paragraph.add_run().add_break(WD_BREAK.LINE)
            for span in line.get("spans", []):
                add_span_to_paragraph(paragraph, span)
    elif fallback_text:
        paragraph.add_run(sanitize_for_xml(str(fallback_text)))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)


def _add_table_event(word_doc, page, event, content_left, available_width, gap_before=0):
    source_table = event["table"]
    if gap_before > 0.5:
        spacer = word_doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.line_spacing = Pt(max(1, gap_before))
        spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        spacer.add_run().font.size = Pt(1)

    rows = max(1, int(source_table.row_count))
    columns = max(1, int(source_table.col_count))
    table = word_doc.add_table(rows=rows, cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    table_width = event["bbox"][2] - event["bbox"][0]
    scale = min(1.0, available_width / max(1, table_width))
    first_row_cells = source_table.rows[0].cells if source_table.rows else []
    column_widths = []
    for column_index in range(columns):
        bbox = first_row_cells[column_index] if column_index < len(first_row_cells) else None
        width = (bbox[2] - bbox[0]) if bbox else table_width / columns
        column_widths.append(max(12, width * scale))

    table_data = source_table.extract()
    for row_index in range(rows):
        source_row = source_table.rows[row_index] if row_index < len(source_table.rows) else None
        row_bboxes = source_row.cells if source_row else []
        source_height = source_row.bbox[3] - source_row.bbox[1] if source_row else 12
        table.rows[row_index].height = Pt(max(10, source_height * scale))
        table.rows[row_index].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            cell.width = Pt(column_widths[column_index])
            bbox = row_bboxes[column_index] if column_index < len(row_bboxes) else None
            fallback = ""
            if row_index < len(table_data) and column_index < len(table_data[row_index]):
                fallback = table_data[row_index][column_index] or ""
            _populate_table_cell(page, cell, bbox, fallback)

    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(round(max(0, event["bbox"][0] - content_left) * 20)))
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    return table


def _add_image_event(word_doc, page, event, content_left, available_width, gap_before=0):
    block = event["block"]
    try:
        img_bbox = fitz.Rect(event["bbox"])
        width_inches = min(img_bbox.width, available_width) / 72.0
        if "image" in block:
            img_bytes = block["image"]
        else:
            pix = page.get_pixmap(clip=img_bbox, dpi=150)
            img_bytes = pix.tobytes("png")
        import io
        word_doc.add_picture(io.BytesIO(img_bytes), width=Inches(width_inches))
        paragraph = word_doc.paragraphs[-1]
        paragraph.paragraph_format.space_before = Pt(max(0, gap_before))
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(max(0, event["bbox"][0] - content_left))
        return paragraph
    except Exception:
        return None


def convert_flow_layout(pdf_doc, word_doc, include_images=True):
    """
    Flow layout: preserve paragraph and table structure while allowing Word to
    reflow text within the page width.
    """
    for page_idx in range(len(pdf_doc)):
        page = pdf_doc[page_idx]
        if page_idx > 0:
            word_doc.add_page_break()
        events = _build_page_events(page, include_images)
        content_left, content_right, _, _ = _page_content_bounds(page, events)
        available_width = page.rect.width - content_left - content_right
        previous_bottom = events[0]["bbox"][1] if events else 0

        for event in events:
            gap_before = max(0, event["bbox"][1] - previous_bottom)
            if event["type"] == "text":
                _add_text_event(word_doc, event, content_left, page.rect.width - content_right, gap_before, exact=False)
            elif event["type"] == "table":
                _add_table_event(word_doc, page, event, content_left, available_width, gap_before)
            elif event["type"] == "image":
                _add_image_event(word_doc, page, event, content_left, available_width, gap_before)
            previous_bottom = max(previous_bottom, event["bbox"][3])


def convert_exact_layout(pdf_doc, word_doc, include_images=True):
    """
    Exact layout: preserve page dimensions, source margins, paragraph line
    breaks and spacing, images, and detected PDF tables.
    """
    for page_idx in range(len(pdf_doc)):
        page = pdf_doc[page_idx]
        page_width = page.rect.width
        page_height = page.rect.height

        section = word_doc.sections[0] if page_idx == 0 else word_doc.add_section(WD_SECTION.NEW_PAGE)
        section.page_width = Emu(int(page_width / 72 * 914400))
        section.page_height = Emu(int(page_height / 72 * 914400))

        events = _build_page_events(page, include_images)
        content_left, content_right, content_top, content_bottom = _page_content_bounds(page, events)
        section.left_margin = Pt(content_left)
        section.right_margin = Pt(content_right)
        section.top_margin = Pt(content_top)
        section.bottom_margin = Pt(content_bottom)
        available_width = page_width - content_left - content_right
        previous_bottom = content_top

        for event in events:
            gap_before = max(0, event["bbox"][1] - previous_bottom)
            if event["type"] == "text":
                _add_text_event(word_doc, event, content_left, page_width - content_right, gap_before, exact=True)
            elif event["type"] == "table":
                _add_table_event(word_doc, page, event, content_left, available_width, gap_before)
            elif event["type"] == "image":
                _add_image_event(word_doc, page, event, content_left, available_width, gap_before)
            previous_bottom = max(previous_bottom, event["bbox"][3])


def perform_ocr_on_page(page):
    """Use Tesseract OCR on a page image to extract text."""
    try:
        import pytesseract
        from PIL import Image
        import io
        
        pix = page.get_pixmap(dpi=300)
        img_bytes = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_bytes))
        
        ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        
        lines = {}
        for i in range(len(ocr_data["text"])):
            text = ocr_data["text"][i].strip()
            if not text:
                continue
            line_num = ocr_data["line_num"][i]
            block_num = ocr_data["block_num"][i]
            key = (block_num, line_num)
            if key not in lines:
                lines[key] = []
            lines[key].append(text)
        
        return [" ".join(words) for words in lines.values()]
    except ImportError:
        return []


def main():
    parser = argparse.ArgumentParser(description="Convert PDF to Word (.docx)")
    parser.add_argument("input_pdf", help="Path to input PDF file")
    parser.add_argument("output_docx", help="Path to output DOCX file")
    parser.add_argument("--layout", choices=["flow", "exact"], default="flow",
                        help="Layout mode: flow (reading order) or exact (preserve positions)")
    parser.add_argument("--images", dest="include_images", action="store_true", default=True,
                        help="Include images in output")
    parser.add_argument("--no-images", dest="include_images", action="store_false",
                        help="Exclude images from output")
    parser.add_argument("--ocr", action="store_true", default=False,
                        help="Use OCR for scanned pages")
    parser.add_argument("--json", action="store_true", default=False,
                        help="Output result as JSON")
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input_pdf):
        result = {"success": False, "error": f"Input file not found: {args.input_pdf}"}
        if args.json:
            print(json.dumps(result))
        else:
            print(f"Error: {result['error']}", file=sys.stderr)
        sys.exit(1)
    
    try:
        with fitz.open(args.input_pdf) as source_pdf:
            total_pages = len(source_pdf)

        warnings = []

        # The dedicated layout engine reconstructs Word sections and columns,
        # which is essential for newsletters, brochures, and other designed
        # PDFs. Keep the in-house converter for flow mode, OCR, image-free
        # exports, and as a safe fallback if the optional engine fails.
        if args.layout == "exact" and args.include_images and not args.ocr:
            try:
                optimization = _convert_with_layout_engine(args.input_pdf, args.output_docx)
                file_size = os.path.getsize(args.output_docx)
                result = {
                    "success": True,
                    "output": args.output_docx,
                    "pages": total_pages,
                    "file_size": file_size,
                    "layout": args.layout,
                    "engine": "pdf2docx",
                    **optimization,
                }
                if args.json:
                    print(json.dumps(result))
                else:
                    print(f"Converted {total_pages} pages to {args.output_docx} ({file_size} bytes)")
                return
            except Exception as layout_error:
                warnings.append(
                    f"Layout engine unavailable; used the built-in converter: {layout_error}"
                )

        pdf_doc = fitz.open(args.input_pdf)
        word_doc = Document()
        
        # Set default font
        style = word_doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Check for scanned/image-only pages if OCR requested
        if args.ocr:
            for page_idx in range(total_pages):
                page = pdf_doc[page_idx]
                text = page.get_text("text").strip()
                
                if len(text) < 50:  # Likely a scanned page
                    ocr_lines = perform_ocr_on_page(page)
                    if ocr_lines:
                        if page_idx > 0:
                            word_doc.add_page_break()
                        for line_text in ocr_lines:
                            paragraph = word_doc.add_paragraph(sanitize_for_xml(line_text))
                            paragraph.paragraph_format.space_after = Pt(2)
                    continue
                
                # Non-scanned page, use normal extraction
                # (handled below — but we need to track which pages we did OCR on)
            
            # Re-open for normal pages 
            pdf_doc.close()
            pdf_doc = fitz.open(args.input_pdf)
        
        # Convert using selected layout
        if args.layout == "flow":
            convert_flow_layout(pdf_doc, word_doc, args.include_images)
        else:
            convert_exact_layout(pdf_doc, word_doc, args.include_images)
        
        # Save
        word_doc.save(args.output_docx)
        pdf_doc.close()
        
        file_size = os.path.getsize(args.output_docx)
        
        result = {
            "success": True,
            "output": args.output_docx,
            "pages": total_pages,
            "file_size": file_size,
            "layout": args.layout,
            "engine": "toolbase",
            "warnings": warnings,
        }
        
        if args.json:
            print(json.dumps(result))
        else:
            print(f"Converted {total_pages} pages to {args.output_docx} ({file_size} bytes)")
    
    except Exception as e:
        result = {"success": False, "error": str(e), "traceback": traceback.format_exc()}
        if args.json:
            print(json.dumps(result))
        else:
            print(f"Error: {e}", file=sys.stderr)
            traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
